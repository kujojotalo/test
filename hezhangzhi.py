from appium import webdriver
from appium.webdriver.extensions.android.nativekey import AndroidKey
import time
import os
from appium.webdriver.common.touch_action import TouchAction
from appium.webdriver.common.multi_action import MultiAction
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import docx
import csv
from PIL import Image
from PIL import ImageChops
import math
from functools import reduce
import operator

doc =docx.Document()
desired_caps = {
  'platformName': 'Android', # 被测手机是安卓
  'platformVersion': '11', # 手机安卓版本
  'deviceName': 'Y66', # 设备名，安卓手机可以随意填写
  'appPackage': 'com.szfish.hzzhi', # 启动APP Package名称
  'appActivity': 'com.szfish.hzzhi.activity.LoginActivity', # 启动Activity名称
  'unicodeKeyboard': True, # 使用自带输入法，输入中文时填True
  'resetKeyboard': True, # 执行完程序恢复原来输入法
  'noReset': True,       # 不要重置App
  'newCommandTimeout': 6000,
  'automationName' : 'UiAutomator2'
  # 'app': r'd:\apk\bili.apk',
}
# 连接Appium Server，初始化自动化环境
driver = webdriver.Remote('http://192.168.1.254:4723/wd/hub', desired_caps)

# 设置缺省等待时间
driver.implicitly_wait(5)

# 根据id定位搜索位置框，点击
#driver.find_element_by_id("et_login_phone").click()
#
# 根据id定位搜索输入框，点击
picture_time = time.strftime("%Y-%m-%d-%H_%M_%S", time.localtime(time.time()))
directory_time = time.strftime("%Y-%m-%d", time.localtime(time.time()))
doc.add_heading('河长制app日常拨测', 0)
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中设置
paragraph =doc.add_paragraph()
directory_time = time.strftime("%Y-%m-%d", time.localtime(time.time()))
run = paragraph.add_run(directory_time)
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中设置
run.bold = True  # 设置字体为粗体
tab = doc.add_table(rows=22, cols=4, style="Table Grid")  # 添加一个4行4列的空表
tab.style.font.size=Pt(12)  # 字体大小
tab.cell(1,1).merge(tab.cell(1,3))# 合并单元格
table_run1 = tab.cell(0,0).paragraphs[0].add_run('测试版本')
table_run1.font.name = u'楷体'
table_run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
table_run1.font.size=Pt(20)  # 字体大小
table_run2 = tab.cell(0,2).paragraphs[0].add_run('测试环境')
table_run2.font.name = u'楷体'
table_run2.element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
table_run2.font.size=Pt(20)  # 字体大小
table_run3 = tab.cell(1,0).paragraphs[0].add_run('登录账号')
table_run3.font.name = u'楷体'
table_run3.element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
table_run3.font.size=Pt(20)  # 字体大小
table_run4 = tab.cell(2,0).paragraphs[0].add_run('系统版本')
table_run4.font.name = u'楷体'
table_run4.element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
table_run4.font.size=Pt(20)  # 字体大小
table_run5 = tab.cell(2,1).paragraphs[0].add_run('测试内容')
table_run5.font.name = u'楷体'
table_run5.element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
table_run5.font.size=Pt(20)  # 字体大小
table_run6 = tab.cell(2,2).paragraphs[0].add_run('测试结果')
table_run6.font.name = u'楷体'
table_run6.element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
table_run6.font.size=Pt(20)  # 字体大小
table_run7 = tab.cell(2,3).paragraphs[0].add_run('备注')
table_run7.font.name = u'楷体'
table_run7.element.rPr.rFonts.set(qn('w:eastAsia'),u'楷体')
table_run7.font.size=Pt(20)  # 字体大小
cell = tab.cell(3, 0)  # 获取某单元格对象（从0开始索引）
cell.text='登录'
cell = tab.cell(4, 0)  # 获取某单元格对象（从0开始索引）
cell.text='新闻动态'
cell = tab.cell(5, 0)  # 获取某单元格对象（从0开始索引）
cell.text='新闻详情'
cell = tab.cell(6, 0)  # 获取某单元格对象（从0开始索引）
cell.text='通知公告'
cell = tab.cell(7, 0)  # 获取某单元格对象（从0开始索引）
cell.text='公告详情'
cell = tab.cell(8, 0)  # 获取某单元格对象（从0开始索引）
cell.text='巡查管理'
cell = tab.cell(9, 0)  # 获取某单元格对象（从0开始索引）
cell.text='巡查查询'
cell = tab.cell(10, 0)  # 获取某单元格对象（从0开始索引）
cell.text='巡查详情'
cell = tab.cell(11, 0)  # 获取某单元格对象（从0开始索引）
cell.text='河道信息'
cell = tab.cell(12, 0)  # 获取某单元格对象（从0开始索引）
cell.text='河道查询'
cell = tab.cell(13, 0)  # 获取某单元格对象（从0开始索引）
cell.text='河道详情'
cell = tab.cell(14, 0)  # 获取某单元格对象（从0开始索引）
cell.text='问题上报'
cell = tab.cell(15, 0)  # 获取某单元格对象（从0开始索引）
cell.text='问题详情'
cell = tab.cell(16, 0)  # 获取某单元格对象（从0开始索引）
cell.text='上报问题'
cell = tab.cell(17, 0)  # 获取某单元格对象（从0开始索引）
cell.text='处理中'
cell = tab.cell(18, 0)  # 获取某单元格对象（从0开始索引）
cell.text='已完成'
cell = tab.cell(19, 0)  # 获取某单元格对象（从0开始索引）
cell.text='监控管理'
cell = tab.cell(20, 0)  # 获取某单元格对象（从0开始索引）
cell.text='监控详情'
cell = tab.cell(21, 0)  # 获取某单元格对象（从0开始索引）
cell.text='公示牌'
cell = tab.cell(3, 1)  # 获取某单元格对象（从0开始索引）
cell.text='用户登录是否正常，首页显示信息是否完整'
cell = tab.cell(4, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(5, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(6, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(7, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(8, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(9, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(10, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(11, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(12, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(13, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(14, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(15, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(16, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(17, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(18, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(19, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(20, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
cell = tab.cell(21, 1)  # 获取某单元格对象（从0开始索引）
cell.text='页面能否正常跳转，能否正常显示'
try:
    File_Path = 'Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app' + '\\' + directory_time + '\\'
    '''File_Path = os.getcwd() + '\\picture\\' + directory_time + '\\' '''
    if not os.path.exists(File_Path):
        os.makedirs(File_Path)
        print("目录新建成功：%s" % File_Path)
    else:
        print("目录已存在！！！")
except BaseException as msg:
    print("新建目录失败：%s" % msg)
class login():
    sbox = driver.find_element_by_id('et_login_phone')
    sbox.send_keys('qhzb')

    sbox = driver.find_element_by_id('et_login_password')
    sbox.send_keys('111111')

    driver.find_element_by_id("btn_login_phone").click()

    time.sleep(1)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '登录' + directory_time + '.png')
    try:
        cell = tab.cell(3, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(3, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'

class news():
    driver.find_element_by_id("tv1").click()

    eles = driver.find_elements_by_id("tv_title")

    for ele in eles:
        # 打印标题
        print(ele.text)
    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '新闻动态' + directory_time + '.png')
    try:
        cell = tab.cell(4, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(4, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("tv_title").click()
    time.sleep(2)

    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '新闻详情' + directory_time + '.png')
    try:
        cell = tab.cell(5, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(5, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    eles = driver.find_elements_by_id("title")
    for ele1 in eles:
        # 打印标题
        print(ele1.text)

    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
class notify():
    driver.find_element_by_id("tv2").click()

    eles = driver.find_elements_by_id("tv_title")

    for ele in eles:
        # 打印标题
        print(ele.text)
    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '通知公告' + directory_time + '.png')
    try:
        cell = tab.cell(6, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(6, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    time.sleep(20)
    driver.find_element_by_id("tv_title").click()
    time.sleep(2)
    eles = driver.find_elements_by_id("title")
    for ele1 in eles:
        # 打印标题
        print(ele1.text)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '公告详情' + directory_time + '.png')
    try:
        cell = tab.cell(7, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(7, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
class inspection():
    driver.find_element_by_id("tv3").click()
    eles = driver.find_elements_by_id("tv_title")

    for ele in eles:
        # 打印标题
        print(ele.text)
    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '巡查管理' + directory_time + '.png')
    try:
        cell = tab.cell(8, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(8, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    time.sleep(2)
    sbox = driver.find_element_by_id('et_content')
    sbox.send_keys('前塘河')
    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '巡查查询' + directory_time + '.png')
    try:
        cell = tab.cell(9, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(9, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    time.sleep(2)
    driver.find_element_by_id("tv_xcjl").click()
    time.sleep(2)
    eles = driver.find_elements_by_id("title")
    for ele1 in eles:
        # 打印标题
        print(ele1.text)
    time.sleep(2)
    driver.find_element_by_id("tv_xc").click()
    time.sleep(2)
    eles = driver.find_elements_by_id("title")
    for ele1 in eles:
        # 打印标题
        print(ele1.text)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '巡查详情' + directory_time + '.png')
    try:
        cell = tab.cell(10, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(10, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
class river():
    driver.find_element_by_id("tv4").click()
    eles = driver.find_elements_by_id("tv_title")
    for ele in eles:
        # 打印标题
        print(ele.text)
    time.sleep(5)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '河道信息' + directory_time + '.png')
    try:
        cell = tab.cell(11, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(11, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    sbox = driver.find_element_by_id('et_content')
    sbox.send_keys('外塘河')
    driver.find_element_by_id("iv_search").click()
    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '河道查询' + directory_time + '.png')
    try:
        cell = tab.cell(12, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(12, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    time.sleep(2)
    driver.find_element_by_id("tv_title").click()
    time.sleep(2)
    eles = driver.find_elements_by_id("title")
    for ele1 in eles:
        # 打印标题
        print(ele1.text)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '河道详情' + directory_time + '.png')
    try:
        cell = tab.cell(13, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(13, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
class problem():
    driver.find_element_by_id("tv5").click()
    time.sleep(2)
    eles = driver.find_elements_by_id("tv_title")
    for ele in eles:
        # 打印标题
        print(ele.text)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '问题上报' + directory_time + '.png')
    try:
        cell = tab.cell(14, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(14, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    time.sleep(2)
    driver.find_element_by_id("tv_title").click()
    time.sleep(5)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '问题详情' + directory_time + '.png')
    try:
        cell = tab.cell(15, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(15, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("next").click()
    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '上报问题' + directory_time + '.png')
    try:
        cell = tab.cell(16, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(16, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
class event():
    driver.find_element_by_id("tv6").click()
    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '处理中' + directory_time + '.png')
    try:
        cell = tab.cell(17, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(17, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_xpath('//android.widget.TextView[@text=\"已完成\"]').click()
    time.sleep(2)
    '''driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '办理中' + directory_time + '.png')
    driver.find_element_by_xpath('//android.widget.TextView[@text=\"已办理\"]').click()
    time.sleep(2)'''
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '已完成' + directory_time + '.png')
    try:
        cell = tab.cell(18, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(18, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("back").click()
    time.sleep(2)

class monitor():
    driver.swipe(500, 500, 500, 300, 500)
    time.sleep(2)
    try:
        driver.find_element_by_id("ll7").click()
    except:
        driver.find_element_by_id("back").click()
        time.sleep(2)
        driver.find_element_by_id("ll7").click()

    time.sleep(2)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '监控管理' + directory_time + '.png')
    try:
        cell = tab.cell(19, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(19, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    # size = driver.get_window_size()
    # print(size)
    # x = size['width']
    # y = size['height']
    # time.sleep(2)

    # 缩放
    def pinch():
        size = driver.get_window_size()
        print(size)
        x = size['width']
        y = size['height']
        time.sleep(2)
        action1 = TouchAction(driver)
        action2 = TouchAction(driver)
        mul_action = MultiAction(driver)
        action1.press(x=x * 0.2, y=y * 0.2).wait(1000).move_to(x=x * 0.2, y=y * 0.2).wait(1000).release()
        action2.press(x=x * 0.8, y=y * 0.8).wait(1000).move_to(x=-x * 0.2, y=-y * 0.2).wait(1000).release()
        # 执行
        print("-----start pinch-----")
        mul_action.add(action1, action2)
        mul_action.perform()

    # 放大
    def zoom():
        size = driver.get_window_size()
        print(size)
        x = size['width']
        y = size['height']
        time.sleep(2)
        action1 = TouchAction(driver)
        action2 = TouchAction(driver)
        mul_action = MultiAction(driver)
        action1.press(x=x * 0.4, y=y * 0.4).wait(1000).move_to(x=-x * 0.2, y=-y * 0.2).wait(1000).release()
        action2.press(x=x * 0.6, y=y * 0.6).wait(1000).move_to(x=x * 0.2, y=y * 0.2).wait(1000).release()
        # 执行
        print("-----start zoom-----")
        mul_action.add(action1, action2)
        mul_action.perform()

    if __name__ == "__main__":

        for i in range(2):
            time.sleep(2)
            zoom()

        for i in range(2):
            time.sleep(2)
            pinch()
    time.sleep(2)
    a1 = 530/850
    b1 = 1080/1920
    x1 = driver.get_window_size()['width']
    y1 = driver.get_window_size()['height']
    TouchAction(driver).press(x=a1*x1, y=b1*y1).release().perform()
    # driver.tap([(364, 545), (365, 575)], 1000)
    time.sleep(2)
    driver.find_element_by_class_name("android.widget.LinearLayout").click()
    time.sleep(10)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '监控详情1' + directory_time + '.png')
    time.sleep(10)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '监控详情2' + directory_time + '.png')

    def image_contrast(img1, img2):
        image1 = Image.open(img1)
        image2 = Image.open(img2)
        h1 = image1.histogram()
        h2 = image2.histogram()
        result = math.sqrt(reduce(operator.add, list(map(lambda a, b: (a - b) ** 2, h1, h2))) / len(h1))
        return result

    img1 = 'Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '监控详情1' + directory_time + '.png'  # 指定图片路径
    img2 = 'Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '监控详情2' + directory_time + '.png'
    result = image_contrast(img1, img2)
    print(result)
    if result>5:
        cell = tab.cell(20, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    else:
        cell = tab.cell(20, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
    driver.find_element_by_id("back").click()
    time.sleep(2)
class billboard():
    driver.find_element_by_id("ll_8").click()
    time.sleep(10)
    driver.get_screenshot_as_file('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '公示牌' + directory_time + '.png')
    try:
        cell = tab.cell(21, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '通过'
    except BaseException:
        cell = tab.cell(21, 2)  # 获取某单元格对象（从0开始索引）
        cell.text = '不通过'
    driver.find_element_by_id("back").click()
    time.sleep(2)
cell = tab.cell(0, 1)  # 获取某单元格对象（从0开始索引）
cell.text='v2.26'
cell = tab.cell(0, 3)  # 获取某单元格对象（从0开始索引）
cell.text='Android7.0'
cell = tab.cell(1, 1)  # 获取某单元格对象（从0开始索引）
cell.text='qhzb、111111'
for r in range(22):#循环将每一行，每一列都设置为居中
    for c in range(4):
        tab.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        tab.cell(r, c).vertical_alignment  = WD_CELL_VERTICAL_ALIGNMENT.CENTER
doc.save('Z:\\02软件测试\\2021年测试文档\\80产品日常拨测\\河长制app\\' + directory_time + '\\' + '河长制app' + directory_time + '.docx')
time.sleep(2)
driver.quit()